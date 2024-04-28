import pandas as pd
from docx import Document
from docx.shared import Cm,Inches,Pt   #Word操作：导入单位换算函数
from docx.enum.text import WD_ALIGN_PARAGRAPH     #导入对齐选项
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import json

def chang_cell(cell):

    paragraph = cell.paragraphs[0]
    run = paragraph.runs[0]
    run.font.size = Pt(10)
    run.font.name = '宋体'
    run.bold = False
    run.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

class Info:
    def __init__(self):
        self.name = "xxx"
        self.tel = '137xx448'
        self.title ='xxxx师'
        self.work = 'xxx职员'
        self.depart = 'xxxxx中心'
        self.company = 'xxxxxxxx公司'

    def read(self,file_name):
        with open(file_name,'r',encoding="utf-8") as f:
            data = json.load(f)
        self.name = data['name']
        self.tel = data['tel']
        self.title = data['title']
        self.work = data['work']
        self.depart = data['depart']
        self.company = data['company']
        self.myname  = data['myname']


def excel_to_word(excel_file, word_file):

    my_info = Info()
    my_info.read("config.json")

    # 读取Excel文件
    excel_data = pd.read_excel(excel_file)

    # 创建Word文档
    doc = Document()
    
    # 添加数据行
    for i, row in excel_data.iterrows():
        

        title = doc.add_paragraph('现职称后专业工作业绩证明')
        title.style.font.name = '宋体'
        title.style.font.size = Pt(24)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1 = title.runs[0]
        run1.bold = True

        # 添加段落
        start_time = str(row[0])[0:10]
        end_time = str(row[1])[0:10]
        achievement = str(row[4])
        role = str(row[5])
        paragraph = doc.add_paragraph('  兹证明'+my_info.company+my_info.depart+my_info.myname+'员工在'+start_time+'至'+end_time+'期间主要取得以下工作业绩：')
        run2 = paragraph.runs[0]
        run2.bold = False
        paragraph.style.font.name = '宋体'
        paragraph.style.font.size = Pt(16)
        table = doc.add_table(rows=2, cols=6)

        cell = table.cell(1, 0)
        cell.text = '1'
        chang_cell(cell)

        cell = table.cell(1, 1)
        cell.width = Inches(1.5)
        cell.text = str(start_time+'至'+end_time)

        chang_cell(cell)


        cell = table.cell(1, 2)
        cell.text = str(achievement)
        cell.width = Inches(2)
        chang_cell(cell)


        cell = table.cell(1, 3)
        cell.text = str(role)
        chang_cell(cell)

        

        cell = table.cell(1, 4)
        cell.text = my_info.depart
        chang_cell(cell)
        
        
        cell = table.cell(1, 5)
        cell.text = my_info.name
        chang_cell(cell)
        
        cell = table.cell(0, 0)
        cell.text = '序号'
        
        chang_cell(cell)

        cell = table.cell(0, 1)
        cell.text = '完成时间'
        chang_cell(cell)
        
        cell = table.cell(0, 2)
        cell.text = '业绩'
        chang_cell(cell)

        cell = table.cell(0, 3)
        cell.text = '角色定位'
        chang_cell(cell)
        
        cell = table.cell(0, 4)
        cell.text = '业务主管部门'
        chang_cell(cell)

        cell = table.cell(0, 5)
        cell.text = '技术负责人'
        chang_cell(cell)
        

        paragraph = doc.add_paragraph('    特此证明。')
        run2 = paragraph.runs[0]
        run2.bold = False
        p1 = doc.add_paragraph('    部门负责人：')
        run2 = p1.runs[0]
        run2.bold = False
        
        p2 = doc.add_paragraph(my_info.depart)
        run2 = p2.runs[0]
        run2.bold = False
        p3 = doc.add_paragraph('二0二四年四月二十日')
        run2 = p3.runs[0]
        run2.bold = False
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p4 = doc.add_paragraph('证明人基本情况介绍')
        run2 = p4.runs[0]
        run2.bold = False
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table = doc.add_table(rows=2, cols=4)

# 设置表格对齐方式为居中对齐
        table.cell(0,0).text = '姓名'
        table.cell(0,1).text = my_info.name
        table.cell(1,0).text = '职称'
        table.cell(1,1).text = my_info.title
        table.cell(0,2).text = '职务'
        table.cell(0,3).text =  my_info.work
        table.cell(1,2).text = '联系电话'
        table.cell(1,3).text =  my_info.tel
        doc.add_page_break()


    # 保存Word文档
    doc.save(word_file)

# 使用示例
import random
# 随机生成一个两位数
random_number = str(random.randint(10, 99))
excel_to_word('后业绩.xlsx', '业绩证明材料'+random_number+'.docx')
